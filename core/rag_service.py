import os
import re
import json
import logging
import hashlib
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass, asdict
from pathlib import Path
import yaml
import time

try:
    import numpy as np
    from sentence_transformers import SentenceTransformer
    import faiss
except ImportError as e:
    print(f"Warning: Required libraries not installed: {e}")
    SentenceTransformer = None
    faiss = None
    np = None

try:
    import fitz  # PyMuPDF
    from docx import Document
    import markdown
except ImportError as e:
    print(f"Warning: Document processing libraries not installed: {e}")
    fitz = None
    Document = None
    markdown = None


@dataclass
class DocumentChunk:
    """Document chunk with metadata"""
    id: str
    content: str
    source: str
    chunk_index: int
    start_char: int
    end_char: int
    metadata: Dict[str, Any]
    embedding: Optional[np.ndarray] = None


@dataclass
class RetrievalResult:
    """RAG retrieval result"""
    chunks: List[DocumentChunk]
    scores: List[float]
    query: str
    total_time: float
    metadata: Dict[str, Any]


class DocumentProcessor:
    """Document processing and loading"""

    def __init__(self, supported_formats: List[str] = None):
        self.supported_formats = supported_formats or ["pdf", "txt", "docx", "md"]
        self.logger = logging.getLogger("document_processor")

    def load_document(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Load document content and metadata"""
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")

        file_extension = file_path.suffix.lower().lstrip('.')

        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported format: {file_extension}")

        metadata = {
            "filename": file_path.name,
            "file_path": str(file_path),
            "file_size": file_path.stat().st_size,
            "format": file_extension,
            "modified_time": file_path.stat().st_mtime
        }

        try:
            if file_extension == "pdf":
                content = self._load_pdf(file_path, metadata)
            elif file_extension == "txt":
                content = self._load_txt(file_path, metadata)
            elif file_extension == "docx":
                content = self._load_docx(file_path, metadata)
            elif file_extension == "md":
                content = self._load_markdown(file_path, metadata)
            else:
                raise ValueError(f"Handler not implemented for: {file_extension}")

            metadata["content_length"] = len(content)
            metadata["word_count"] = len(content.split())

            self.logger.info(f"Loaded document: {file_path.name} ({metadata['word_count']} words)")
            return content, metadata

        except Exception as e:
            self.logger.error(f"Failed to load document {file_path}: {e}")
            raise

    def _load_pdf(self, file_path: Path, metadata: Dict[str, Any]) -> str:
        """Load PDF using PyMuPDF"""
        if fitz is None:
            raise ImportError("PyMuPDF not installed")

        try:
            doc = fitz.open(file_path)
            content_parts = []

            metadata["page_count"] = len(doc)
            metadata["pdf_metadata"] = doc.metadata

            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()

                if text.strip():
                    content_parts.append(f"\n--- Page {page_num + 1} ---\n")
                    content_parts.append(text)

            doc.close()
            full_content = "\n".join(content_parts)
            return self._clean_text(full_content)

        except Exception as e:
            self.logger.error(f"Error reading PDF {file_path}: {e}")
            raise

    def _load_txt(self, file_path: Path, metadata: Dict[str, Any]) -> str:
        """Load text file"""
        try:
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    metadata["encoding"] = encoding
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise UnicodeDecodeError("Could not decode file with any encoding")

            return self._clean_text(content)

        except Exception as e:
            self.logger.error(f"Error reading TXT {file_path}: {e}")
            raise

    def _load_docx(self, file_path: Path, metadata: Dict[str, Any]) -> str:
        """Load DOCX file"""
        if Document is None:
            raise ImportError("python-docx not installed")

        try:
            doc = Document(file_path)
            content_parts = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        content_parts.append(row_text)

            metadata["paragraph_count"] = len(doc.paragraphs)
            metadata["table_count"] = len(doc.tables)

            full_content = "\n".join(content_parts)
            return self._clean_text(full_content)

        except Exception as e:
            self.logger.error(f"Error reading DOCX {file_path}: {e}")
            raise

    def _load_markdown(self, file_path: Path, metadata: Dict[str, Any]) -> str:
        """Load Markdown file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            # Convert markdown to plain text if markdown library available
            if markdown:
                html_content = markdown.markdown(content)
                import re
                text_content = re.sub(r'<[^>]+>', '', html_content)
                content = text_content

            return self._clean_text(content)

        except Exception as e:
            self.logger.error(f"Error reading Markdown {file_path}: {e}")
            raise

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)

        # Normalize line breaks
        text = re.sub(r'\n\s*\n', '\n\n', text)

        # Strip lines
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(line for line in lines if line)

        return text.strip()


class TextChunker:
    """Text chunking for document processing"""

    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.logger = logging.getLogger("text_chunker")

    def chunk_document(self, content: str, source: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Split document into chunks"""
        chunks = []

        # Split into paragraphs first
        paragraphs = self._split_into_paragraphs(content)

        current_chunk = ""
        current_start = 0
        chunk_index = 0

        for paragraph in paragraphs:
            if len(paragraph) > self.chunk_size:
                # Save current chunk if exists
                if current_chunk.strip():
                    chunks.append(self._create_chunk(
                        current_chunk, source, chunk_index,
                        current_start, current_start + len(current_chunk), metadata
                    ))
                    chunk_index += 1

                # Split long paragraph
                sub_chunks = self._split_long_paragraph(paragraph)
                for sub_chunk in sub_chunks:
                    start_pos = content.find(sub_chunk, current_start)
                    chunks.append(self._create_chunk(
                        sub_chunk, source, chunk_index,
                        start_pos, start_pos + len(sub_chunk), metadata
                    ))
                    chunk_index += 1

                current_chunk = ""
                current_start = content.find(paragraph, current_start) + len(paragraph)

            else:
                potential_chunk = current_chunk + "\n" + paragraph if current_chunk else paragraph

                if len(potential_chunk) <= self.chunk_size:
                    current_chunk = potential_chunk
                else:
                    # Save current chunk
                    if current_chunk.strip():
                        chunks.append(self._create_chunk(
                            current_chunk, source, chunk_index,
                            current_start, current_start + len(current_chunk), metadata
                        ))
                        chunk_index += 1

                    # Start new chunk with overlap
                    if self.chunk_overlap > 0 and current_chunk:
                        overlap_text = current_chunk[-self.chunk_overlap:]
                        current_chunk = overlap_text + "\n" + paragraph
                        current_start = current_start + len(current_chunk) - len(overlap_text) - len(paragraph) - 1
                    else:
                        current_chunk = paragraph
                        current_start = content.find(paragraph, current_start)

        # Save final chunk
        if current_chunk.strip():
            chunks.append(self._create_chunk(
                current_chunk, source, chunk_index,
                current_start, current_start + len(current_chunk), metadata
            ))

        self.logger.info(f"Created {len(chunks)} chunks for document: {source}")
        return chunks

    def _split_into_paragraphs(self, content: str) -> List[str]:
        """Split content into paragraphs"""
        paragraphs = re.split(r'\n\s*\n', content)
        return [p.strip() for p in paragraphs if p.strip()]

    def _split_long_paragraph(self, paragraph: str) -> List[str]:
        """Split overly long paragraphs"""
        sentences = re.split(r'(?<=[.!?])\s+', paragraph)

        chunks = []
        current_chunk = ""

        for sentence in sentences:
            potential_chunk = current_chunk + " " + sentence if current_chunk else sentence

            if len(potential_chunk) <= self.chunk_size:
                current_chunk = potential_chunk
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence

        if current_chunk:
            chunks.append(current_chunk.strip())

        return chunks

    def _create_chunk(self, content: str, source: str, chunk_index: int,
                      start_char: int, end_char: int, doc_metadata: Dict[str, Any]) -> DocumentChunk:
        """Create document chunk"""
        chunk_id = self._generate_chunk_id(source, chunk_index)

        chunk_metadata = {
            **doc_metadata,
            "chunk_method": "paragraph_based",
            "chunk_size_config": self.chunk_size,
            "chunk_overlap_config": self.chunk_overlap
        }

        return DocumentChunk(
            id=chunk_id,
            content=content.strip(),
            source=source,
            chunk_index=chunk_index,
            start_char=start_char,
            end_char=end_char,
            metadata=chunk_metadata
        )

    def _generate_chunk_id(self, source: str, chunk_index: int) -> str:
        """Generate unique chunk ID"""
        source_hash = hashlib.md5(source.encode()).hexdigest()[:8]
        return f"{source_hash}_{chunk_index:04d}"


class VectorStore:
    """Vector storage and retrieval"""

    def __init__(self, embedding_model: SentenceTransformer, index_type: str = "IndexFlatIP"):
        self.embedding_model = embedding_model
        self.index_type = index_type
        self.index: Optional[faiss.Index] = None
        self.chunks: List[DocumentChunk] = []
        self.id_to_index: Dict[str, int] = {}
        self.logger = logging.getLogger("vector_store")

    def add_chunks(self, chunks: List[DocumentChunk]):
        """Add chunks to vector store"""
        if not chunks:
            return

        self.logger.info(f"Adding {len(chunks)} chunks to vector store...")

        # Generate embeddings
        texts = [chunk.content for chunk in chunks]
        embeddings = self.embedding_model.encode(texts, show_progress_bar=True)

        # Add embeddings to chunks
        for chunk, embedding in zip(chunks, embeddings):
            chunk.embedding = embedding

        # Create index if needed
        if self.index is None:
            self._create_index(embeddings)

        # Add to FAISS index
        start_index = len(self.chunks)
        self.index.add(embeddings)

        # Update storage
        for i, chunk in enumerate(chunks):
            self.chunks.append(chunk)
            self.id_to_index[chunk.id] = start_index + i

        self.logger.info(f"Vector store now contains {len(self.chunks)} chunks")

    def _create_index(self, embeddings: np.ndarray):
        """Create FAISS index"""
        dimension = embeddings.shape[1]

        if self.index_type == "IndexFlatIP":
            self.index = faiss.IndexFlatIP(dimension)
        elif self.index_type == "IndexFlatL2":
            self.index = faiss.IndexFlatL2(dimension)
        else:
            self.index = faiss.IndexFlatIP(dimension)

        self.logger.info(f"Created FAISS index: {self.index_type} with dimension {dimension}")

    def search(self, query: str, top_k: int = 5, threshold: float = 0.0) -> RetrievalResult:
        """Search for relevant chunks"""
        start_time = time.time()

        if self.index is None or len(self.chunks) == 0:
            self.logger.warning("Vector store is empty")
            return RetrievalResult([], [], query, 0.0, {"warning": "empty_store"})

        # Encode query
        query_embedding = self.embedding_model.encode([query])

        # Search
        scores, indices = self.index.search(query_embedding, min(top_k, len(self.chunks)))

        # Process results
        result_chunks = []
        result_scores = []

        for score, idx in zip(scores[0], indices[0]):
            if idx != -1 and score >= threshold:
                result_chunks.append(self.chunks[idx])
                result_scores.append(float(score))

        total_time = time.time() - start_time

        metadata = {
            "total_chunks_searched": len(self.chunks),
            "threshold_applied": threshold,
            "query_length": len(query)
        }

        self.logger.debug(f"Search completed in {total_time:.3f}s, found {len(result_chunks)} results")

        return RetrievalResult(
            chunks=result_chunks,
            scores=result_scores,
            query=query,
            total_time=total_time,
            metadata=metadata
        )

    def save(self, save_path: str):
        """Save vector store"""
        save_path = Path(save_path)
        save_path.mkdir(parents=True, exist_ok=True)

        # Save FAISS index
        if self.index:
            index_path = save_path / "faiss_index.bin"
            faiss.write_index(self.index, str(index_path))

        # Save chunks (without embeddings)
        chunks_data = []
        for chunk in self.chunks:
            chunk_dict = asdict(chunk)
            chunk_dict.pop('embedding', None)
            chunks_data.append(chunk_dict)

        chunks_path = save_path / "chunks.json"
        with open(chunks_path, 'w', encoding='utf-8') as f:
            json.dump(chunks_data, f, ensure_ascii=False, indent=2)

        # Save ID mapping
        mapping_path = save_path / "id_mapping.json"
        with open(mapping_path, 'w', encoding='utf-8') as f:
            json.dump(self.id_to_index, f, indent=2)

        self.logger.info(f"Vector store saved to {save_path}")

    def load(self, load_path: str):
        """Load vector store"""
        load_path = Path(load_path)

        if not load_path.exists():
            raise FileNotFoundError(f"Vector store not found: {load_path}")

        # Load FAISS index
        index_path = load_path / "faiss_index.bin"
        if index_path.exists():
            self.index = faiss.read_index(str(index_path))

        # Load chunks
        chunks_path = load_path / "chunks.json"
        if chunks_path.exists():
            with open(chunks_path, 'r', encoding='utf-8') as f:
                chunks_data = json.load(f)

            self.chunks = []
            for chunk_dict in chunks_data:
                chunk = DocumentChunk(**chunk_dict)
                self.chunks.append(chunk)

        # Load ID mapping
        mapping_path = load_path / "id_mapping.json"
        if mapping_path.exists():
            with open(mapping_path, 'r', encoding='utf-8') as f:
                self.id_to_index = json.load(f)

        self.logger.info(f"Vector store loaded from {load_path}")


class RAGService:
    """Main RAG service"""

    def __init__(self, config_path: str = "config.yaml"):
        self.logger = logging.getLogger("rag_service")
        self.config = self._load_config(config_path)

        # Initialize components
        self.document_processor = DocumentProcessor(
            self.config.get("rag.document_formats", ["pdf", "txt", "docx", "md"])
        )

        self.text_chunker = TextChunker(
            chunk_size=self.config.get("embedding.chunk_size", 500),
            chunk_overlap=self.config.get("embedding.chunk_overlap", 50)
        )

        # Initialize embedding model
        self.embedding_model = self._load_embedding_model()

        # Initialize vector store
        self.vector_store = VectorStore(
            self.embedding_model,
            self.config.get("rag.index_type", "IndexFlatIP")
        )

        # Try to load existing index
        self._try_load_existing_index()

        self.logger.info("RAG service initialized successfully")

    def _load_config(self, config_path: str) -> Dict[str, Any]:
        """Load configuration"""
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                return yaml.safe_load(f)
        except Exception as e:
            self.logger.error(f"Failed to load config: {e}")
            return {}

    def _load_embedding_model(self) -> SentenceTransformer:
        """Load embedding model"""
        if SentenceTransformer is None:
            raise ImportError("sentence-transformers not installed")

        model_name = self.config.get("embedding.model_name", "BAAI/bge-small-en-v1.5")
        cache_folder = self.config.get("embedding.cache_folder", "./data/models/embeddings")
        device = self.config.get("embedding.device", "cpu")

        self.logger.info(f"Loading embedding model: {model_name}")

        try:
            model = SentenceTransformer(
                model_name,
                cache_folder=cache_folder,
                device=device
            )

            self.logger.info(f"Embedding model loaded successfully on {device}")
            return model

        except Exception as e:
            self.logger.error(f"Failed to load embedding model: {e}")
            raise

    def _try_load_existing_index(self):
        """Try to load existing vector index"""
        vector_db_path = self.config.get("rag.vector_db_path", "./data/vectors")

        try:
            self.vector_store.load(vector_db_path)
            self.logger.info("Loaded existing vector index")
        except (FileNotFoundError, Exception):
            self.logger.info("No existing vector index found, will create new one")

    def add_documents_from_directory(self, directory_path: str, save_index: bool = True) -> int:
        """Add all documents from directory"""
        directory_path = Path(directory_path)

        if not directory_path.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")

        # Find all supported documents
        supported_extensions = self.document_processor.supported_formats

        document_paths = []
        for ext in supported_extensions:
            paths = list(directory_path.glob(f"*.{ext}"))
            document_paths.extend([str(p) for p in paths])

        self.logger.info(f"Found {len(document_paths)} documents in {directory_path}")

        return self.add_documents(document_paths, save_index)

    def add_documents(self, document_paths: List[str], save_index: bool = True) -> int:
        """Add documents to RAG system"""
        total_chunks = 0

        for doc_path in document_paths:
            try:
                # Load document
                content, metadata = self.document_processor.load_document(doc_path)

                # Create chunks
                chunks = self.text_chunker.chunk_document(content, doc_path, metadata)

                # Add to vector store
                self.vector_store.add_chunks(chunks)

                total_chunks += len(chunks)
                self.logger.info(f"Added {len(chunks)} chunks from {doc_path}")

            except Exception as e:
                self.logger.error(f"Failed to process document {doc_path}: {e}")
                continue

        # Save index
        if save_index and total_chunks > 0:
            self.save_index()

        self.logger.info(f"Total chunks added: {total_chunks}")
        return total_chunks

    def search(self, query: str, top_k: int = None, threshold: float = None) -> RetrievalResult:
        """Search for relevant documents"""
        if top_k is None:
            top_k = self.config.get("rag.top_k", 5)

        if threshold is None:
            threshold = self.config.get("rag.similarity_threshold", 0.0)

        return self.vector_store.search(query, top_k, threshold)

    def get_context_for_query(self, query: str, max_length: int = None) -> str:
        """Get formatted context for query"""
        if max_length is None:
            max_length = self.config.get("rag.max_context_length", 2000)

        # Retrieve relevant documents
        retrieval_result = self.search(query)

        if not retrieval_result.chunks:
            return "No relevant policy documents found for this query."

        # Build context
        context_parts = []
        current_length = 0

        for i, (chunk, score) in enumerate(zip(retrieval_result.chunks, retrieval_result.scores)):
            chunk_text = f"**Source {i + 1}** (Score: {score:.3f}):\n{chunk.content}\n"

            if current_length + len(chunk_text) > max_length:
                break

            context_parts.append(chunk_text)
            current_length += len(chunk_text)

        if not context_parts:
            return "No relevant policy documents found for this query."

        context = "\n".join(context_parts)
        header = f"Retrieved {len(context_parts)} relevant policy sections:\n\n"

        return header + context

    def save_index(self):
        """Save vector index"""
        vector_db_path = self.config.get("rag.vector_db_path", "./data/vectors")
        self.vector_store.save(vector_db_path)

    def get_statistics(self) -> Dict[str, Any]:
        """Get RAG system statistics"""
        return {
            "total_chunks": len(self.vector_store.chunks),
            "unique_sources": len(set(chunk.source for chunk in self.vector_store.chunks)),
            "embedding_model": self.config.get("embedding.model_name"),
            "index_type": self.config.get("rag.index_type"),
            "chunk_size": self.config.get("embedding.chunk_size"),
            "chunk_overlap": self.config.get("embedding.chunk_overlap")
        }

    def rebuild_index(self, document_directory: str = None):
        """Rebuild vector index"""
        self.logger.info("Rebuilding vector index...")

        # Clear current index
        self.vector_store = VectorStore(
            self.embedding_model,
            self.config.get("rag.index_type", "IndexFlatIP")
        )

        # Re-add documents
        if document_directory:
            doc_dir = document_directory
        else:
            doc_dir = self.config.get("paths.documents_dir", "./data/documents")

        if os.path.exists(doc_dir):
            self.add_documents_from_directory(doc_dir, save_index=True)
        else:
            self.logger.warning(f"Document directory not found: {doc_dir}")


def create_rag_service(config_path: str = "config.yaml") -> RAGService:
    """Create RAG service instance"""
    return RAGService(config_path)